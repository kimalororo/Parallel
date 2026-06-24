"""Restyle the SGA report body to match the provided coursework example.

The title page is intentionally left untouched: the script only changes
paragraphs and tables after the first manual page break.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
SOURCE_DOCX = ROOT / "sga_parallel_report.docx"
OUTPUT_DOCX = ROOT / "sga_parallel_report_like_example.docx"


def set_run_font(run, *, name="Times New Roman", size=12, bold=None, italic=None, all_caps=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if all_caps is not None:
        run.font.all_caps = all_caps


def has_page_break(paragraph) -> bool:
    for br in paragraph._p.xpath(".//w:br"):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def first_body_paragraph_index(doc: Document) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if has_page_break(paragraph):
            return index + 1
    return 0


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def get_or_create_style(doc: Document, name: str, base_style: str = "Normal"):
    styles = doc.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles[base_style]
    return style


def configure_paragraph_style(
    style,
    *,
    font_name: str,
    size: float,
    bold: bool | None = None,
    italic: bool | None = None,
    all_caps: bool | None = None,
    align=None,
    first_line_indent=None,
    left_indent=None,
    right_indent=None,
    space_before=0,
    space_after=0,
    line_spacing=1.5,
    line_spacing_rule=None,
):
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if all_caps is not None:
        font.all_caps = all_caps

    pf = style.paragraph_format
    pf.alignment = align
    pf.first_line_indent = first_line_indent
    pf.left_indent = left_indent
    pf.right_indent = right_indent
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if line_spacing_rule is not None:
        pf.line_spacing_rule = line_spacing_rule


def ensure_example_styles(doc: Document) -> None:
    styles = {
        "Основной текст1": dict(
            font_name="Times New Roman",
            size=12,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=Cm(1.25),
            line_spacing=1.5,
        ),
        "Заголовок отчёта": dict(
            font_name="Times New Roman",
            size=12,
            bold=True,
            align=None,
            line_spacing=1.5,
        ),
        "*Раздел основной части": dict(
            font_name="Times New Roman",
            size=12,
            bold=True,
            all_caps=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.5,
        ),
        "*Подраздел основной части": dict(
            font_name="Times New Roman",
            size=12,
            bold=True,
            align=None,
            line_spacing=1.5,
        ),
        "Подпись к рисунку": dict(
            font_name="Times New Roman",
            size=12,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.25,
        ),
        "Подпись к таблице": dict(
            font_name="Times New Roman",
            size=12,
            align=None,
            line_spacing=1.25,
        ),
        "Элемент списка": dict(
            font_name="Times New Roman",
            size=12,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line_spacing=1.5,
        ),
        "Листинг кода": dict(
            font_name="Courier New",
            size=12,
            align=None,
            line_spacing=1.0,
        ),
    }

    for name, options in styles.items():
        style = get_or_create_style(doc, name)
        configure_paragraph_style(style, **options)


def normalize_caption(text: str, kind: str, number: int) -> str:
    pattern = rf"^{kind}\s*\d+\s*[.—-]\s*(.*)$"
    match = re.match(pattern, text.strip())
    tail = match.group(1).strip() if match else text.strip()
    return f"{kind} {number} — {tail}"


def remove_cell_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for shading in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(shading)


def style_body_paragraph(paragraph) -> None:
    text = paragraph.text.strip()
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    if not text and not paragraph.runs:
        return

    if text.startswith("Рисунок"):
        paragraph.style = "Подпись к рисунку"
        return

    if text.startswith("Таблица"):
        paragraph.style = "Подпись к таблице"
        return

    if not text and paragraph.runs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    if text == "Ключевые сведения":
        paragraph.style = "Заголовок отчёта"
        return

    if re.match(r"^\d+\.\d+", text):
        paragraph.style = "*Подраздел основной части"
        return

    if re.match(r"^\d+\.\s", text):
        paragraph.style = "*Раздел основной части"
        return

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.style = "Основной текст1"


def style_body_table(table) -> None:
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            remove_cell_shading(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, size=11, bold=(row_index == 0))


def restyle() -> Path:
    shutil.copyfile(SOURCE_DOCX, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)
    ensure_example_styles(doc)
    body_start = first_body_paragraph_index(doc)

    table_no = 1
    figure_no = 1
    for paragraph in doc.paragraphs[body_start:]:
        text = paragraph.text.strip()
        if text.startswith("Таблица"):
            replace_paragraph_text(paragraph, normalize_caption(text, "Таблица", table_no))
            table_no += 1
        elif text.startswith("Рисунок"):
            replace_paragraph_text(paragraph, normalize_caption(text, "Рисунок", figure_no))
            figure_no += 1
        style_body_paragraph(paragraph)

    # Table 0 belongs to the title page and must remain exactly as it was.
    for table in doc.tables[1:]:
        style_body_table(table)

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(restyle())
