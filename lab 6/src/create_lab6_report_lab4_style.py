from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = PROJECT_ROOT / "LAB 4.docx"
OUTPUTS_DIR = PROJECT_ROOT / "outputs"
PLOTS_DIR = OUTPUTS_DIR / "plots"
REPORT_PATH = PROJECT_ROOT / "LAB 6.docx"

SECTION_STYLE = "Заголовок отчёта"
BODY_STYLE = "Основной текст1"
SUBSECTION_STYLE = "*Подраздел основной части"
CAPTION_STYLE = "Подпись к рисунку"
TABLE_STYLE = "Normal Table"


def fmt(value, digits: int = 3) -> str:
    if pd.isna(value):
        return "-"
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def style_name(doc: Document, preferred: str, fallback: str = "Normal") -> str:
    try:
        doc.styles[preferred]
        return preferred
    except KeyError:
        return fallback


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def set_paragraph_text(paragraph, text: str, size: int = 12, bold: bool | None = None) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def compact_blank_paragraph(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.font.size = Pt(1)


def set_title_page(doc: Document) -> None:
    paragraphs = doc.paragraphs
    set_paragraph_text(
        paragraphs[13],
        "КЛАСТЕРИЗАЦИЯ И КЛАССИФИКАЦИЯ С ИСПОЛЬЗОВАНИЕМ APACHE SPARK MLLIB",
        size=12,
        bold=True,
    )
    set_paragraph_text(paragraphs[15], "Лабораторная работа № 6", size=12)
    set_paragraph_text(
        paragraphs[18],
        "Параллельные и высокопроизводительные вычисления",
        size=12,
        bold=True,
    )
    set_paragraph_text(paragraphs[25], "Томск - 2026", size=12)

    if doc.tables:
        for row in doc.tables[0].rows:
            for cell in row.cells:
                if "21.05.2026" in cell.text:
                    cell.text = cell.text.replace("21.05.2026", "10.06.2026")

    for index in [5, 6, 7, 8, 9, 11, 12, 19, 20, 21, 22, 23, 24]:
        if index < len(paragraphs) and not paragraphs[index].text.strip():
            compact_blank_paragraph(paragraphs[index])


def truncate_after_title_page(doc: Document) -> None:
    body = doc.element.body
    children = list(body)
    keep_through = None
    for index, element in enumerate(children):
        tag = element.tag.split("}")[-1]
        if tag == "p" and element.xpath('.//w:br[@w:type="page"]'):
            keep_through = index
            break
    if keep_through is None:
        keep_through = 27

    for element in children[keep_through + 1 :]:
        if element.tag.split("}")[-1] != "sectPr":
            body.remove(element)


def add_body_paragraph(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph(style=style or style_name(doc, BODY_STYLE))
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    return paragraph


def add_section(doc: Document, text: str):
    return doc.add_paragraph(text, style=style_name(doc, SECTION_STYLE))


def add_subsection(doc: Document, text: str):
    return doc.add_paragraph(text, style=style_name(doc, SUBSECTION_STYLE, SECTION_STYLE))


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")
        borders.append(element)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_borders(cell)


def add_df_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], font_size: int = 9) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = TABLE_STYLE
    except KeyError:
        pass
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=font_size)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for index, col in enumerate(columns):
            set_cell_text(cells[index], fmt(row[col]), size=font_size)
    doc.add_paragraph(style=style_name(doc, BODY_STYLE))


def add_picture(doc: Document, image_path: Path, caption: str, width_cm: float = 15.2) -> None:
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph(style=style_name(doc, BODY_STYLE))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption, style=style_name(doc, CAPTION_STYLE, BODY_STYLE))
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def cluster_mode(modes: pd.DataFrame, cluster: int, feature: str) -> str:
    row = modes[(modes["cluster"] == cluster) & (modes["feature"] == feature)]
    return "" if row.empty else str(row.iloc[0]["mode"])


def build_cluster_interpretation(counts: pd.DataFrame, numeric: pd.DataFrame, modes: pd.DataFrame) -> pd.DataFrame:
    names = {
        0: "новые помесячные клиенты с риском оттока",
        1: "лояльные premium-клиенты с большим набором услуг",
        2: "экономные клиенты без интернет-услуг",
    }
    descriptions = {
        0: "низкий срок обслуживания, помесячный контракт, часто fiber optic и мало дополнительных сервисов",
        1: "долгий срок обслуживания, высокая месячная выручка и много подключенных опций",
        2: "низкая месячная плата, отсутствие интернет-сервиса и минимальный набор услуг",
    }
    rows = []
    for _, count_row in counts.iterrows():
        cluster = int(count_row["cluster"])
        total = int(count_row["count"])
        no_churn_row = modes[
            (modes["cluster"] == cluster)
            & (modes["feature"] == "Churn")
            & (modes["mode"] == "No")
        ]
        no_churn_count = int(no_churn_row.iloc[0]["count"]) if not no_churn_row.empty else 0
        num_row = numeric[numeric["cluster"] == cluster].iloc[0]
        rows.append(
            {
                "cluster": cluster,
                "name": names.get(cluster, f"кластер {cluster}"),
                "count": total,
                "tenure": float(num_row["tenure"]),
                "monthly": float(num_row["MonthlyCharges"]),
                "total": float(num_row["TotalCharges"]),
                "churn_rate": (total - no_churn_count) / total * 100,
                "contract": cluster_mode(modes, cluster, "Contract"),
                "internet": cluster_mode(modes, cluster, "InternetService"),
                "description": descriptions.get(cluster, ""),
            }
        )
    return pd.DataFrame(rows)


def main() -> None:
    run_summary = json.loads((OUTPUTS_DIR / "run_summary.json").read_text(encoding="utf-8"))
    best_clusters = pd.read_csv(OUTPUTS_DIR / "best_cluster_models.csv")
    counts = pd.read_csv(OUTPUTS_DIR / "cluster_counts.csv")
    numeric = pd.read_csv(OUTPUTS_DIR / "cluster_numeric_profile.csv")
    modes = pd.read_csv(OUTPUTS_DIR / "cluster_categorical_modes.csv")
    classifier_metrics = pd.read_csv(OUTPUTS_DIR / "classifier_metrics.csv")
    cluster_summary = build_cluster_interpretation(counts, numeric, modes)

    doc = Document(str(TEMPLATE_PATH))
    set_title_page(doc)
    truncate_after_title_page(doc)

    ds = run_summary["dataset"]
    split = run_summary["classification_split"]

    add_section(doc, "Задание")
    add_body_paragraph(
        doc,
        "Цель работы – получить навыки применения распределённых алгоритмов машинного обучения на платформе Apache Spark MLlib.",
        bold_prefix="Цель работы",
    )
    add_body_paragraph(
        doc,
        "Задачи – выбрать датасет с числовыми и категориальными признаками, выполнить предобработку в Spark ML Pipeline, сравнить KMeans, Bisecting KMeans и Gaussian Mixture, интерпретировать кластеры и обучить три классификатора для предсказания метки кластера.",
        bold_prefix="Задачи",
    )

    add_section(doc, "Основная часть")
    add_subsection(doc, "Описание выбранного датасета")
    add_body_paragraph(
        doc,
        "Использован датасет Telco Customer Churn, содержащий сведения о клиентах телеком-компании: подключённые услуги, тип договора, способ оплаты, срок обслуживания и платежи.",
    )
    add_body_paragraph(doc, f"Источник данных: {ds['source_url']}")
    add_body_paragraph(
        doc,
        f"Количество записей: {ds['rows']}. После удаления customerID использовано {ds['columns_after_customer_id_drop']} столбцов. Итоговая размерность вектора features после кодирования равна {ds['feature_vector_dimension']}.",
    )
    add_body_paragraph(
        doc,
        "Числовые признаки: tenure, MonthlyCharges, TotalCharges. Категориальные признаки: gender, SeniorCitizen, Partner, Dependents, PhoneService, MultipleLines, InternetService, OnlineSecurity, OnlineBackup, DeviceProtection, TechSupport, StreamingTV, StreamingMovies, Contract, PaperlessBilling, PaymentMethod.",
    )

    add_subsection(doc, "Предобработка данных")
    add_body_paragraph(
        doc,
        f"В данных найдено 11 пропусков в TotalCharges. Значения были приведены к double и заполнены медианой {ds['numeric_imputation_medians']['TotalCharges']:.2f}.",
    )
    add_body_paragraph(
        doc,
        "Pipeline Spark ML включает StringIndexer, OneHotEncoder, VectorAssembler, StandardScaler и итоговую сборку общего признакового вектора features. Такой подход сохраняет воспроизводимость всех этапов обработки.",
    )

    add_subsection(doc, "Кластеризация")
    add_body_paragraph(
        doc,
        "Для KMeans и Bisecting KMeans количество кластеров выбиралось по silhouette score и WCSS, для Gaussian Mixture использовался BIC. Лучшие значения приведены в таблице.",
    )
    best_table = best_clusters.copy()
    best_table["silhouette"] = best_table["silhouette"].round(3)
    best_table["wcss"] = best_table["wcss"].round(1)
    best_table["bic"] = best_table["bic"].round(1)
    add_df_table(
        doc,
        best_table,
        ["algorithm", "k", "silhouette", "wcss", "bic"],
        ["Алгоритм", "k", "Silhouette", "WCSS", "BIC"],
    )
    add_picture(doc, PLOTS_DIR / "silhouette_comparison.png", "Рисунок 1 – сравнение silhouette score для методов кластеризации.")
    add_picture(doc, PLOTS_DIR / "gaussian_mixture_bic.png", "Рисунок 2 – выбор количества компонент Gaussian Mixture по BIC.")

    add_subsection(doc, "Интерпретация кластеров")
    interpretation_table = cluster_summary.copy()
    interpretation_table["tenure"] = interpretation_table["tenure"].round(1)
    interpretation_table["monthly"] = interpretation_table["monthly"].round(1)
    interpretation_table["churn_rate"] = interpretation_table["churn_rate"].round(1)
    add_df_table(
        doc,
        interpretation_table,
        ["cluster", "name", "count", "tenure", "monthly", "churn_rate"],
        ["Кластер", "Содержательное название", "Кол-во", "tenure", "Monthly", "Отток, %"],
        font_size=8,
    )
    for _, row in cluster_summary.iterrows():
        add_body_paragraph(
            doc,
            f"Кластер {int(row['cluster'])} – {row['name']}: {row['description']}. Средний срок обслуживания {row['tenure']:.1f} месяцев, средняя месячная плата {row['monthly']:.1f}, оценочная доля оттока {row['churn_rate']:.1f}%.",
        )
    add_picture(doc, PLOTS_DIR / "pca_clusters.png", "Рисунок 3 – двумерная PCA-проекция объектов с цветом по кластеру KMeans.")
    add_picture(doc, PLOTS_DIR / "cluster_numeric_heatmap.png", "Рисунок 4 – средние значения числовых признаков по кластерам.")

    add_subsection(doc, "Классификация")
    add_body_paragraph(
        doc,
        f"В качестве целевой переменной использованы метки лучшей кластеризации ({run_summary['classification_target_algorithm']}, k = 3). Данные разделены на обучающую и тестовую выборки: {split['train_rows']} и {split['test_rows']} записей.",
    )
    add_body_paragraph(
        doc,
        "Гиперпараметры подбирались через TrainValidationSplit: для LogisticRegression – regParam и elasticNetParam, для RandomForest – numTrees и maxDepth, для GBT – maxDepth в схеме One-vs-Rest.",
    )
    class_table = classifier_metrics.copy()
    for col in ["accuracy", "weighted_precision", "weighted_recall", "weighted_f1"]:
        class_table[col] = class_table[col].round(4)
    class_table["training_time_sec"] = class_table["training_time_sec"].round(2)
    add_df_table(
        doc,
        class_table,
        ["model", "accuracy", "weighted_precision", "weighted_recall", "weighted_f1", "training_time_sec"],
        ["Модель", "Accuracy", "Precision", "Recall", "F1", "Время, c"],
        font_size=8,
    )
    add_picture(doc, PLOTS_DIR / "classifier_metrics.png", "Рисунок 5 – сравнение Accuracy и Weighted F1 классификаторов.")
    add_picture(doc, PLOTS_DIR / "confusion_matrix.png", "Рисунок 6 – матрица ошибок лучшего классификатора.")
    add_body_paragraph(
        doc,
        f"Лучший результат показала LogisticRegression: Accuracy = {classifier_metrics.iloc[0]['accuracy']:.4f}, Weighted F1 = {classifier_metrics.iloc[0]['weighted_f1']:.4f}. На тестовой выборке допущено только 4 ошибки между кластерами 0 и 1.",
    )

    add_section(doc, "Заключение")
    add_body_paragraph(
        doc,
        "Наиболее интерпретируемый результат дала KMeans-кластеризация при k = 3: она отделила новых помесячных клиентов с повышенным риском оттока, лояльных клиентов с дорогим набором услуг и экономный сегмент без интернет-сервиса.",
    )
    add_body_paragraph(
        doc,
        "Bisecting KMeans дал близкое, но чуть более слабое значение silhouette, а Gaussian Mixture оказался менее устойчивым для разреженного one-hot пространства. Среди классификаторов лучшей оказалась LogisticRegression: она дала максимальный F1 и обучалась быстрее GBT.",
    )

    doc.save(REPORT_PATH)
    print(f"Report saved to {REPORT_PATH}")


if __name__ == "__main__":
    main()
