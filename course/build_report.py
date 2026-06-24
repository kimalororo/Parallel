"""Build the final DOCX report for topic 1."""

from __future__ import annotations

import csv
import importlib.util
from pathlib import Path
from statistics import mean

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
RESULTS_DIR = ROOT / "results"
PLOTS_DIR = RESULTS_DIR / "plots"
OUTPUT_DOCX = ROOT / "sga_parallel_report.docx"
TABLE_GEOM_PATH = Path(
    r"C:\Users\kimal\.codex\plugins\cache\openai-primary-runtime\documents\26.622.11653\skills\documents\scripts\table_geometry.py"
)


def load_table_geometry():
    spec = importlib.util.spec_from_file_location("table_geometry", TABLE_GEOM_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load table geometry helper from {TABLE_GEOM_PATH}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module.apply_table_geometry


apply_table_geometry = load_table_geometry()


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as file:
        return list(csv.DictReader(file))


def parse_float(value: str) -> float:
    return float(value.replace(",", "."))


def set_run_font(run, *, name="Calibri", size=11, color="000000", bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph(paragraph, *, before=0, after=6, line=1.1, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.alignment = align


def set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string("2E74B5")
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.10

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string("2E74B5")
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.10

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string("1F4D78")
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.10


def set_section_geometry(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def add_paragraph(doc: Document, text: str, *, size=11, bold=False, italic=False, color="000000",
                  align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.1, style=None):
    p = doc.add_paragraph(style=style)
    set_paragraph(p, before=before, after=after, line=line, align=align)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_spacer(doc: Document, before: int, after: int = 0) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=before, after=after, line=1.0)


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def style_table(table, header_fill: str = "F2F4F7") -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
        if row_index == 0:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), header_fill)
                tc_pr.append(shd)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def set_cell_text(cell, text: str, *, size=10, bold=False, italic=False, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=4, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_run_font(run, size=10, bold=True)


def add_figure(doc: Document, image_path: Path, caption: str, width=6.3) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=6, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    set_paragraph(cap, before=2, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=10, italic=True, color="434343")


def add_title_page(doc: Document) -> None:
    add_spacer(doc, 36)
    title = add_paragraph(
        doc,
        "Отчет",
        size=26,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
    )
    subtitle = add_paragraph(
        doc,
        "по теме 1: параллельная реализация простого генетического алгоритма (SGA) для оптимизации математических функций",
        size=14,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=10,
    )
    summary = add_paragraph(
        doc,
        "В работе реализованы последовательная и параллельная версии SGA, проведены эксперименты на функциях Sphere, Rastrigin и Rosenbrock, а также измерены время выполнения, ускорение и эффективность параллелизации.",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=14,
    )

    meta_title = doc.add_paragraph()
    set_paragraph(meta_title, before=4, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    meta_run = meta_title.add_run("Ключевые сведения")
    set_run_font(meta_run, size=12, bold=True, color="1F4D78")

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    rows = [
        ("Тема", "Параллельная SGA для оптимизации математических функций"),
        ("Среда", "Windows 11, Python 3.12.13, 16 логических ядер"),
        ("Метод", "Tournament selection, arithmetic crossover, gaussian mutation"),
        ("Эксперименты", "9 конфигураций сходимости, 5 повторов для speedup, eval_repeats = 50"),
        ("Результат", "Максимальное измеренное ускорение составило 3.35x на 8 процессах"),
    ]
    for i, (left, right) in enumerate(rows):
        set_cell_text(table.cell(i, 0), left, size=10, bold=True, color="1F1F1F")
        set_cell_text(table.cell(i, 1), right, size=10)
    apply_table_geometry(
        table,
        [1800, 7560],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )


def add_environment_table(doc: Document) -> None:
    add_table_caption(doc, "Таблица 1. Среда выполнения и параметры эксперимента")
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows = [
        ("Параметр", "Значение"),
        ("Операционная система", "Windows 11"),
        ("Интерпретатор", "Python 3.12.13"),
        ("Число логических ядер", "16"),
        ("Библиотека параллелизации", "multiprocessing.Pool"),
        ("Повторов для speedup", "5"),
    ]
    for i, (left, right) in enumerate(rows):
        set_cell_text(table.cell(i, 0), left, size=10, bold=(i == 0))
        set_cell_text(table.cell(i, 1), right, size=10, bold=(i == 0))
    apply_table_geometry(
        table,
        [2200, 7160],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    style_table(table, header_fill="F2F4F7")


def add_benchmark_table(doc: Document) -> None:
    add_table_caption(doc, "Таблица 2. Тестовые функции, использованные в исследовании")
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows = [
        ("Функция", "Формула и минимум", "Назначение"),
        ("Sphere", "f(x) = Σ x_i^2; f* = 0", "Базовая выпуклая функция без локальных ловушек"),
        (
            "Rastrigin",
            "f(x) = 10D + Σ[x_i^2 - 10 cos(2πx_i)]; f* = 0",
            "Многомодальная функция для проверки способности к глобальному поиску",
        ),
        (
            "Rosenbrock",
            "f(x) = Σ[100(x_{i+1} - x_i^2)^2 + (1 - x_i)^2]; f* = 0",
            "Классическая узкая долина для проверки устойчивости сходимости",
        ),
    ]
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, size=9.5, bold=(i == 0))
    apply_table_geometry(
        table,
        [1400, 3600, 4360],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    style_table(table, header_fill="E8EEF5")


def add_convergence_table(doc: Document) -> None:
    rows = [row for row in read_csv_rows(RESULTS_DIR / "experiment_runs.csv") if row["label"] == "convergence_grid"]
    rows.sort(key=lambda row: (int(row["population_size"]), int(row["generations"])))
    add_table_caption(doc, "Таблица 3. Сходимость SGA на функции Rastrigin при разных размерах популяции и числе поколений")
    table = doc.add_table(rows=10, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Популяция", "Поколения", "Лучшее значение", "Время, с"]
    for j, header in enumerate(headers):
        set_cell_text(table.cell(0, j), header, size=9.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows, start=1):
        set_cell_text(table.cell(i, 0), row["population_size"], size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(i, 1), row["generations"], size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(i, 2), f"{parse_float(row['best_value']):.6f}", size=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(table.cell(i, 3), f"{parse_float(row['runtime_sec']):.3f}", size=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
    apply_table_geometry(
        table,
        [1600, 1600, 3100, 3060],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    style_table(table, header_fill="F2F4F7")


def add_speedup_table(doc: Document) -> None:
    rows = read_csv_rows(RESULTS_DIR / "speedup_summary.csv")
    add_table_caption(doc, "Таблица 4. Ускорение и эффективность параллельной версии")
    table = doc.add_table(rows=5, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Процессы", "Повторов", "Среднее время, с", "Ускорение", "Эффективность, %"]
    for j, header in enumerate(headers):
        set_cell_text(table.cell(0, j), header, size=9.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows, start=1):
        set_cell_text(table.cell(i, 0), row["processes"], size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(i, 1), row["runs"], size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(i, 2), f"{parse_float(row['mean_runtime_sec']):.3f}", size=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(table.cell(i, 3), f"{parse_float(row['speedup']):.2f}", size=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(table.cell(i, 4), f"{parse_float(row['efficiency_pct']):.1f}", size=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
    apply_table_geometry(
        table,
        [1300, 1100, 2400, 1700, 2860],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    style_table(table, header_fill="E8EEF5")


def add_quality_table(doc: Document) -> None:
    rows = [row for row in read_csv_rows(RESULTS_DIR / "experiment_runs.csv") if row["label"] == "function_quality"]
    add_table_caption(doc, "Таблица 5. Итоговое качество на трех тестовых функциях")
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Функция", "Лучшее значение", "Время, с"]
    for j, header in enumerate(headers):
        set_cell_text(table.cell(0, j), header, size=9.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows, start=1):
        set_cell_text(table.cell(i, 0), row["objective"], size=9.5)
        set_cell_text(table.cell(i, 1), f"{parse_float(row['best_value']):.6f}", size=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(table.cell(i, 2), f"{parse_float(row['runtime_sec']):.3f}", size=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
    apply_table_geometry(
        table,
        [2200, 3000, 4160],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    style_table(table, header_fill="F2F4F7")


def add_project_files_table(doc: Document) -> None:
    add_table_caption(doc, "Таблица 6. Состав проекта")
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows = [
        ("Файл", "Назначение"),
        ("sga_parallel.py", "Последовательная и параллельная реализация SGA"),
        ("run_experiments.py", "Запуск экспериментов, сохранение CSV и генерация графиков"),
        ("regenerate_plots.py", "Повторная отрисовка графиков из сохраненных CSV"),
        ("results/", "Итоговые таблицы и графики для отчета"),
    ]
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, size=9.5, bold=(i == 0))
    apply_table_geometry(
        table,
        [2200, 7160],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    style_table(table, header_fill="F2F4F7")


def add_body_sections(doc: Document) -> None:
    heading(doc, "1. Введение", level=1)
    add_paragraph(
        doc,
        "Эволюционные алгоритмы хорошо подходят для задач оптимизации, в которых целевая функция сложна, многомодальна или вычислительно дорога. В этой работе исследуется простой генетический алгоритм (Simple Genetic Algorithm, SGA), а основное внимание уделяется тому, как распараллеливание оценки приспособленности влияет на общее время выполнения и качество поиска.",
    )
    add_paragraph(
        doc,
        "Практическая цель работы состоит в том, чтобы реализовать последовательную и параллельную версии SGA, провести сравнимые эксперименты на стандартных тестовых функциях и сделать вывод о том, при каких условиях параллельная реализация дает ощутимый выигрыш.",
    )

    heading(doc, "2. Постановка задачи", level=1)
    add_paragraph(
        doc,
        "Для каждой особи генетический алгоритм хранит вещественный вектор размерности D. На каждом поколении выполняются селекция, кроссинговер, мутация и оценка приспособленности. В последовательной версии все вычисления происходят в одном процессе; в параллельной версии наиболее дорогая операция, то есть оценка фитнеса, распределяется между несколькими процессами через multiprocessing.Pool.",
    )
    add_paragraph(
        doc,
        "В экспериментальной части сравниваются: время одного запуска, среднее время на поколение, итоговое лучшее значение целевой функции, ускорение Speedup(p) = T(1) / T(p) и эффективность Efficiency(p) = Speedup(p) / p × 100%.",
    )

    heading(doc, "3. Теоретические основы", level=1)
    heading(doc, "3.1. Простая генетическая схема", level=2)
    add_paragraph(
        doc,
        "SGA моделирует естественный отбор в дискретной популяции. На каждом шаге из текущего множества решений выбираются родители, между ними выполняется кроссинговер, затем потомки мутируют, и после этого в новое поколение попадают лучшие решения по значению целевой функции. В данной работе используется турнирная селекция, арифметический кроссинговер и гауссова мутация.",
    )
    add_paragraph(
        doc,
        "Такой вариант алгоритма удобен тем, что не требует сложных структур представления и хорошо подходит для непрерывной оптимизации. При этом он достаточно универсален, чтобы сравнивать влияние размеров популяции и числа поколений на качество и скорость сходимости.",
    )

    heading(doc, "3.2. Почему параллелится именно фитнес", level=2)
    add_paragraph(
        doc,
        "Самая затратная часть SGA обычно связана не с самой селекцией, а с вычислением значения целевой функции для каждой особи. Эти вычисления независимы друг от друга, поэтому их можно распределять по процессам без сложной синхронизации. Именно этот тип нагрузки и используют как типичный пример embarrassingly parallel задачи.",
    )
    add_paragraph(
        doc,
        "Чтобы увидеть выигрыш отчетливее, в эксперименте фитнес дополнительно утяжелен повторным вычислением одной и той же функции eval_repeats = 50 раз. Это не меняет саму постановку оптимизации, но позволяет честнее сравнить последовательную и параллельную схему по времени.",
    )

    heading(doc, "3.3. Тестовые функции", level=2)
    add_paragraph(
        doc,
        "Для оценки поведения алгоритма выбраны три стандартные функции. Sphere показывает поведение на гладкой выпуклой поверхности, Rastrigin проверяет способность искать глобальный минимум среди множества локальных, а Rosenbrock демонстрирует устойчивость на узкой искривленной долине.",
    )
    add_benchmark_table(doc)

    heading(doc, "4. Реализация", level=1)
    heading(doc, "4.1. Структура проекта", level=2)
    add_paragraph(
        doc,
        "Реализация разделена на три небольших файла. В sga_parallel.py находится ядро алгоритма и функции оптимизации; run_experiments.py запускает пакет экспериментов, сохраняет результаты в CSV и строит графики; regenerate_plots.py позволяет восстановить изображения из уже сохраненных данных без повторного долгого расчета.",
    )
    add_project_files_table(doc)

    heading(doc, "4.2. Последовательная и параллельная версии", level=2)
    add_paragraph(
        doc,
        "Обе версии используют общий поток логики: случайная инициализация популяции, оценка фитнеса, выбор родителей методом турнира, арифметический кроссинговер и гауссову мутацию. Элитарность сохранена на уровне двух лучших особей, чтобы устойчиво удерживать найденные хорошие решения.",
    )
    add_paragraph(
        doc,
        "Разница между версиями заключается только в вычислении fitness. В последовательном варианте кандидат оценивается обычным циклом, а в параллельном варианте список особей передается в Pool.map. Такое разбиение минимально вмешивается в алгоритм и хорошо отражает типичный рабочий сценарий для задач оценки моделей, функций или симуляций.",
    )
    add_paragraph(
        doc,
        "Для запусков использовались популяция 300 и 400 поколений в speedup-блоке, а также девять конфигураций сходимости на функции Rastrigin: размеры популяции 50, 100 и 500 при 100, 500 и 1000 поколениях.",
    )

    heading(doc, "4.3. Параметры алгоритма", level=2)
    add_paragraph(
        doc,
        "В базовом варианте использованы турнир размера 3, вероятность кроссинговера 0.85, вероятность мутации около 1 / D и среднеквадратичная величина мутационного шага 8% от диапазона поиска. Эти значения не подбирались агрессивно под одну функцию, а выбраны как разумный компромисс для общего экспериментального сравнения.",
    )
    add_environment_table(doc)

    heading(doc, "5. Экспериментальные результаты", level=1)
    heading(doc, "5.1. Сходимость на Rastrigin", level=2)
    add_paragraph(
        doc,
        "Ниже приведена полная таблица для девяти конфигураций. Видно, что увеличение числа поколений в целом улучшает решение, но конкретный один запуск остается стохастическим и не обязан вести себя строго монотонно. Поэтому вывод делается по совокупности конфигураций, а не по одной точке.",
    )
    add_convergence_table(doc)
    add_figure(
        doc,
        PLOTS_DIR / "convergence.png",
        "Рисунок 1. Сходимость SGA на функции Rastrigin для нескольких репрезентативных конфигураций",
    )

    heading(doc, "5.2. Ускорение и эффективность", level=2)
    add_paragraph(
        doc,
        "Среднее ускорение растет с числом процессов, но эффективность падает, потому что накладные расходы на создание процессов, обмен данными и синхронизацию начинают занимать заметную долю времени. Это нормальный и ожидаемый результат для стандартной модели multiprocessing на ограниченной задаче.",
    )
    add_speedup_table(doc)
    add_figure(
        doc,
        PLOTS_DIR / "speedup.png",
        "Рисунок 2. Ускорение параллельной версии при росте числа процессов",
    )

    heading(doc, "5.3. Сравнение качества на трех тестовых функциях", level=2)
    add_paragraph(
        doc,
        "Функция Sphere решается почти до нуля, что ожидаемо для гладкой и простой поверхности. На Rastrigin качество заметно хуже, поскольку функция содержит множество локальных минимумов. Rosenbrock остается самой сложной из трех, и это хорошо видно по итоговому значению фитнеса.",
    )
    add_quality_table(doc)
    add_figure(
        doc,
        PLOTS_DIR / "function_quality.png",
        "Рисунок 3. Итоговое качество на трех тестовых функциях",
    )

    heading(doc, "6. Обсуждение результатов", level=1)
    add_paragraph(
        doc,
        "Эксперименты подтверждают, что параллелизация оценки фитнеса оправдана тогда, когда сама функция достаточно дорогая. При eval_repeats = 50 выигрыш становится отчетливым уже на четырех процессах, а на восьми процессах алгоритм ускоряется более чем в три раза по сравнению с последовательной версией.",
    )
    add_paragraph(
        doc,
        "При этом параллельная версия не превращается в линейный ускоритель. Ограничения накладывают накладные расходы Python multiprocessing, время на сериализацию данных, а также то, что не вся работа в SGA распараллеливается. Поэтому эффективность постепенно снижается от 88.6% на двух процессах до 41.8% на восьми.",
    )
    add_paragraph(
        doc,
        "По качеству поиска наиболее убедительный результат дала конфигурация с небольшой популяцией и большим числом поколений. Это согласуется с интуицией: если дать алгоритму достаточно итераций, он способен глубже исследовать ландшафт функции, но итоговая точность все равно зависит от случайности выбора родителей и мутаций.",
    )

    heading(doc, "7. Заключение", level=1)
    add_paragraph(
        doc,
        "В работе реализован простой генетический алгоритм для оптимизации непрерывных функций и его параллельная версия на multiprocessing.Pool. Экспериментальная часть показала, что параллельная оценка фитнеса дает реальный выигрыш во времени, особенно при дорогой функции, но этот выигрыш не является линейным и зависит от баланса между вычислениями и накладными расходами.",
    )
    add_paragraph(
        doc,
        "Практический вывод состоит в том, что SGA удобно использовать как базовую схему для параллелизации, а дальнейшее улучшение качества можно получить за счет более тщательной настройки параметров, увеличения числа повторов и более сложных стратегий адаптации мутаций и селекции.",
    )

    heading(doc, "8. Список литературы", level=1)
    references = [
        "1. Holland J. H. Adaptation in Natural and Artificial Systems. University of Michigan Press, 1975.",
        "2. Goldberg D. E. Genetic Algorithms in Search, Optimization, and Machine Learning. Addison-Wesley, 1989.",
        "3. Python Software Foundation. multiprocessing - Process-based parallelism. https://docs.python.org/3/library/multiprocessing.html",
        "4. Rastrigin L. A. Systems of extremal control. 1974.",
        "5. Rosenbrock H. H. An Automatic Method for Finding the Greatest or Least Value of a Function. The Computer Journal, 1960.",
    ]
    for ref in references:
        add_paragraph(doc, ref, size=10.5, after=4)

    heading(doc, "Приложение A. Состав поставки", level=1)
    add_paragraph(
        doc,
        "Ниже перечислены основные артефакты проекта, которые позволяют воспроизвести исследование и при необходимости повторно построить графики или отчет.",
    )
    add_project_files_table(doc)


def build_report() -> Path:
    doc = Document()
    set_styles(doc)
    section = doc.sections[0]
    set_section_geometry(section)
    add_title_page(doc)
    doc.add_page_break()
    add_body_sections(doc)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    path = build_report()
    print(f"Saved {path}")
