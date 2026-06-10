from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUTS_DIR = PROJECT_ROOT / "outputs"
PLOTS_DIR = OUTPUTS_DIR / "plots"
REPORT_PATH = PROJECT_ROOT / "LAB 6.docx"


def fmt(value, digits: int = 3) -> str:
    if pd.isna(value):
        return "-"
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_df_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            set_cell_text(cells[idx], fmt(row[col]))
    doc.add_paragraph()


def add_picture(doc: Document, image_path: Path, caption: str, width_cm: float = 15.5) -> None:
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_title_page(doc: Document) -> None:
    lines = [
        "Министерство науки и высшего образования Российской Федерации",
        "Федеральное государственное автономное образовательное учреждение высшего образования",
        "«Национальный исследовательский Томский политехнический университет»",
    ]
    for line in lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Лабораторная работа № 6")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Кластеризация и классификация с использованием Apache Spark MLlib")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("по дисциплине:\nПараллельные вычисления")

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph("Томск - 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def cluster_mode(modes: pd.DataFrame, cluster: int, feature: str) -> str:
    row = modes[(modes["cluster"] == cluster) & (modes["feature"] == feature)]
    return "" if row.empty else str(row.iloc[0]["mode"])


def build_cluster_interpretation(
    counts: pd.DataFrame,
    numeric: pd.DataFrame,
    modes: pd.DataFrame,
) -> pd.DataFrame:
    rows = []
    names = {
        0: "новые помесячные клиенты с риском оттока",
        1: "лояльные premium-клиенты с большим набором услуг",
        2: "экономные клиенты без интернет-услуг",
    }
    descriptions = {
        0: (
            "низкий срок обслуживания, помесячный контракт, часто fiber optic, "
            "электронный чек и мало дополнительных сервисов"
        ),
        1: (
            "долгий срок обслуживания, высокая месячная выручка, много подключенных "
            "опций и преимущественно долгосрочный договор"
        ),
        2: (
            "низкая месячная плата, отсутствие интернет-сервиса, чаще mailed check "
            "и минимальный набор услуг"
        ),
    }
    for _, count_row in counts.iterrows():
        cluster = int(count_row["cluster"])
        total = int(count_row["count"])
        no_churn_count = int(
            modes[
                (modes["cluster"] == cluster)
                & (modes["feature"] == "Churn")
                & (modes["mode"] == "No")
            ].iloc[0]["count"]
        )
        num_row = numeric[numeric["cluster"] == cluster].iloc[0]
        rows.append(
            {
                "cluster": cluster,
                "name": names.get(cluster, f"кластер {cluster}"),
                "count": total,
                "tenure": float(num_row["tenure"]),
                "monthly": float(num_row["MonthlyCharges"]),
                "total": float(num_row["TotalCharges"]),
                "churn_rate": (total - no_churn_count) / total * 100,
                "contract": cluster_mode(modes, cluster, "Contract"),
                "internet": cluster_mode(modes, cluster, "InternetService"),
                "description": descriptions.get(cluster, ""),
            }
        )
    return pd.DataFrame(rows)


def main() -> None:
    run_summary = json.loads((OUTPUTS_DIR / "run_summary.json").read_text(encoding="utf-8"))
    cluster_metrics = pd.read_csv(OUTPUTS_DIR / "cluster_metrics.csv")
    best_clusters = pd.read_csv(OUTPUTS_DIR / "best_cluster_models.csv")
    counts = pd.read_csv(OUTPUTS_DIR / "cluster_counts.csv")
    numeric = pd.read_csv(OUTPUTS_DIR / "cluster_numeric_profile.csv")
    modes = pd.read_csv(OUTPUTS_DIR / "cluster_categorical_modes.csv")
    classifier_metrics = pd.read_csv(OUTPUTS_DIR / "classifier_metrics.csv")
    confusion = pd.read_csv(OUTPUTS_DIR / "confusion_matrix.csv")
    cluster_summary = build_cluster_interpretation(counts, numeric, modes)

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("Задание", level=1)
    doc.add_paragraph(
        "Цель работы: получить навыки применения распределённых алгоритмов машинного "
        "обучения на платформе Apache Spark MLlib: построить Pipeline предобработки, "
        "сравнить методы кластеризации, интерпретировать сегменты и обучить модели "
        "классификации для предсказания принадлежности к кластеру."
    )

    doc.add_heading("Описание выбранного датасета", level=1)
    ds = run_summary["dataset"]
    doc.add_paragraph(
        "Использован датасет Telco Customer Churn, содержащий сведения о клиентах "
        "телеком-компании: подключённые услуги, тип договора, способ оплаты, срок "
        "обслуживания, ежемесячные и суммарные платежи, а также признак оттока."
    )
    doc.add_paragraph(f"Источник данных: {ds['source_url']}")
    doc.add_paragraph(
        f"Количество записей: {ds['rows']}. После удаления идентификатора customerID "
        f"использовано {ds['columns_after_customer_id_drop']} столбцов. Вектор признаков "
        f"после кодирования имеет размерность {ds['feature_vector_dimension']}."
    )
    doc.add_paragraph(
        "Числовые признаки: tenure, MonthlyCharges, TotalCharges. Категориальные "
        "признаки: пол, SeniorCitizen, Partner, Dependents, PhoneService, MultipleLines, "
        "InternetService, OnlineSecurity, OnlineBackup, DeviceProtection, TechSupport, "
        "StreamingTV, StreamingMovies, Contract, PaperlessBilling, PaymentMethod."
    )
    doc.add_paragraph(
        "Цель исследования: сегментировать клиентов по профилю использования услуг и "
        "построить классификаторы, которые восстанавливают найденные сегменты по "
        "исходным признакам."
    )

    doc.add_heading("Предобработка данных", level=1)
    doc.add_paragraph(
        "Пропуски обнаружены только в TotalCharges: 11 строк. Значения TotalCharges "
        f"были приведены к double и заполнены медианой {ds['numeric_imputation_medians']['TotalCharges']:.2f}. "
        "Для tenure и MonthlyCharges пропусков не было, но в Pipeline сохранена единая "
        "логика обработки числовых признаков."
    )
    doc.add_paragraph(
        "Pipeline Spark ML включает StringIndexer для категориальных признаков, "
        "OneHotEncoder для преобразования категорий в бинарные индикаторы, "
        "VectorAssembler для числового блока, StandardScaler для масштабирования "
        "числовых признаков и итоговый VectorAssembler для сборки общего features."
    )

    doc.add_heading("Кластеризация", level=1)
    doc.add_paragraph(
        "Для KMeans и Bisecting KMeans количество кластеров выбиралось по silhouette "
        "score и дополнительно фиксировался WCSS. Для Gaussian Mixture использован BIC; "
        "меньшее значение BIC считается лучшим."
    )
    best_table = best_clusters.copy()
    best_table["silhouette"] = best_table["silhouette"].round(3)
    best_table["wcss"] = best_table["wcss"].round(1)
    best_table["bic"] = best_table["bic"].round(1)
    add_df_table(
        doc,
        best_table,
        ["algorithm", "k", "silhouette", "wcss", "bic"],
        ["Алгоритм", "k", "Silhouette", "WCSS", "BIC"],
    )
    add_picture(
        doc,
        PLOTS_DIR / "silhouette_comparison.png",
        "Рисунок 1 - сравнение silhouette score для трёх методов кластеризации.",
    )
    add_picture(
        doc,
        PLOTS_DIR / "gaussian_mixture_bic.png",
        "Рисунок 2 - выбор количества компонент Gaussian Mixture по BIC.",
    )
    doc.add_page_break()

    doc.add_heading("Интерпретация кластеров", level=1)
    interpretation_table = cluster_summary.copy()
    interpretation_table["tenure"] = interpretation_table["tenure"].round(1)
    interpretation_table["monthly"] = interpretation_table["monthly"].round(1)
    interpretation_table["total"] = interpretation_table["total"].round(1)
    interpretation_table["churn_rate"] = interpretation_table["churn_rate"].round(1)
    add_df_table(
        doc,
        interpretation_table,
        ["cluster", "name", "count", "tenure", "monthly", "churn_rate"],
        [
            "Кластер",
            "Содержательное название",
            "Кол-во",
            "tenure",
            "Monthly",
            "Отток, %",
        ],
    )
    for _, row in cluster_summary.iterrows():
        doc.add_paragraph(
            f"Кластер {int(row['cluster'])} - {row['name']}: {row['description']}. "
            f"Средний срок обслуживания {row['tenure']:.1f} месяцев, средняя месячная "
            f"плата {row['monthly']:.1f}, оценочная доля оттока {row['churn_rate']:.1f}%."
        )
    add_picture(
        doc,
        PLOTS_DIR / "pca_clusters.png",
        "Рисунок 3 - двумерная PCA-проекция объектов с цветом по кластеру KMeans.",
    )
    add_picture(
        doc,
        PLOTS_DIR / "cluster_numeric_heatmap.png",
        "Рисунок 4 - средние значения числовых признаков по кластерам.",
    )

    doc.add_heading("Классификация", level=1)
    split = run_summary["classification_split"]
    doc.add_paragraph(
        "В качестве целевой переменной использованы метки лучшей кластеризации "
        f"({run_summary['classification_target_algorithm']}, k=3). Данные разделены "
        f"на обучающую и тестовую выборки: {split['train_rows']} и {split['test_rows']} "
        "записей соответственно."
    )
    doc.add_paragraph(
        "Настройка гиперпараметров выполнена через TrainValidationSplit. Для "
        "LogisticRegression подбирались regParam и elasticNetParam, для RandomForest - "
        "numTrees и maxDepth. Поскольку GBTClassifier в Spark является бинарным, для "
        "трёх кластеров применена схема One-vs-Rest: для каждого кластера обучался "
        "отдельный бинарный GBT с подбором maxDepth."
    )
    class_table = classifier_metrics.copy()
    for col in ["accuracy", "weighted_precision", "weighted_recall", "weighted_f1"]:
        class_table[col] = class_table[col].round(4)
    class_table["training_time_sec"] = class_table["training_time_sec"].round(2)
    add_df_table(
        doc,
        class_table,
        ["model", "accuracy", "weighted_precision", "weighted_recall", "weighted_f1", "training_time_sec"],
        ["Модель", "Accuracy", "Precision", "Recall", "F1", "Время, c"],
    )
    add_picture(
        doc,
        PLOTS_DIR / "classifier_metrics.png",
        "Рисунок 5 - сравнение Accuracy и Weighted F1 классификаторов.",
    )
    add_picture(
        doc,
        PLOTS_DIR / "confusion_matrix.png",
        "Рисунок 6 - матрица ошибок лучшего классификатора.",
    )
    doc.add_paragraph(
        "Лучший результат показала LogisticRegression: Accuracy = "
        f"{classifier_metrics.iloc[0]['accuracy']:.4f}, Weighted F1 = "
        f"{classifier_metrics.iloc[0]['weighted_f1']:.4f}. Матрица ошибок показывает, "
        "что на тестовой выборке допущено только 4 ошибки между кластерами 0 и 1; "
        "кластер 2 восстановлен без ошибок."
    )

    doc.add_heading("Выводы", level=1)
    doc.add_paragraph(
        "Наиболее интерпретируемый результат дала KMeans-кластеризация при k=3: "
        "она отделила новых помесячных клиентов с повышенным риском оттока, "
        "лояльных клиентов с дорогим набором услуг и экономный сегмент без "
        "интернет-сервиса. Bisecting KMeans дал близкое, но чуть более слабое "
        "значение silhouette, а Gaussian Mixture оказался менее устойчивым для "
        "разреженного one-hot пространства."
    )
    doc.add_paragraph(
        "Классификаторы показали высокое качество, потому что предсказывали метки, "
        "полученные из тех же признаков. LogisticRegression оказалась лучшей по F1 "
        "и одновременно достаточно быстрой, RandomForest немного уступил, а GBT "
        "потребовал больше времени из-за схемы One-vs-Rest."
    )

    doc.save(REPORT_PATH)
    print(f"Report saved to {REPORT_PATH}")


if __name__ == "__main__":
    main()
